"""Genera Minority Report DOCX usando Azure AI Foundry + plantilla XOC.

La IA devuelve JSON; este script genera el DOCX final.
"""

from __future__ import annotations

import argparse
import json
import re
import zipfile
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from dotenv import load_dotenv

from azure_minority_client import DEFAULT_PAYLOAD_PATH, generate_minority_payload

ROOT = Path(__file__).resolve().parent
TEMPLATE_PATH = ROOT / "Plantilla Minority Report.docx"
MODEL_TEMPLATE_PATH = ROOT / "TXDX-WR-2606-26083-JOCKEY MINORITY_REPORT_SEMANAL_26_JUNIO.docx"
OUTPUT_DIR = ROOT / "output"

BRAND_GREEN = RGBColor(0x00, 0xFF, 0x9F)
BRAND_BLUE = RGBColor(0x00, 0xF0, 0xFF)
MODEL_BLUE = RGBColor(0x00, 0x6D, 0x9F)
DARK_TEXT = RGBColor(0x1F, 0x38, 0x62)
MUTED_TEXT = RGBColor(0x5C, 0x66, 0x70)
TABLE_HEADER = "0B1F2A"
TABLE_ALT = "EAF8F5"
CALLOUT = "EAF8F5"
WARNING = "FFF4D6"
BORDER = "8FDACC"


def base_template_path() -> Path:
    if not TEMPLATE_PATH.is_file():
        raise FileNotFoundError(f"No existe la plantilla base de Minority Report: {TEMPLATE_PATH}")
    return TEMPLATE_PATH


def _safe_name(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9_-]+", "-", value.upper().strip())
    return value.strip("-")[:80] or "CLIENTE"


def _output_path(payload: dict[str, Any]) -> Path:
    document_code = _clean_text(payload.get("document_code"))
    if document_code:
        return OUTPUT_DIR / f"{_safe_name(document_code)}.docx"
    client = _safe_name(payload.get("client_name") or "CLIENTE")
    period = _safe_name(payload.get("period") or date.today().isoformat())
    return OUTPUT_DIR / f"MINORITY-REPORT-XOC_{client}_{period}.docx"


def _writable_output_path(path: Path) -> Path:
    if not path.exists():
        return path
    try:
        with path.open("r+b"):
            return path
    except PermissionError:
        for version in range(2, 100):
            candidate = path.with_name(f"{path.stem}_v{version}{path.suffix}")
            if not candidate.exists():
                return candidate
        raise


def _as_list(value: Any) -> list[Any]:
    if isinstance(value, list):
        return value
    return [value] if value not in (None, "") else []


def _clean_text(value: Any) -> str:
    return str(value or "").strip()


def _textbox_text(box: Any) -> str:
    return "".join(node.text or "" for node in box.iter(qn("w:t")))


def _set_textbox_text(box: Any, value: str) -> None:
    nodes = list(box.iter(qn("w:t")))
    if not nodes:
        return
    nodes[0].text = value
    for node in nodes[1:]:
        node.text = ""


def _replace_textbox_jockey_client(box: Any, client: str) -> None:
    nodes = list(box.iter(qn("w:t")))
    for index, node in enumerate(nodes):
        if (node.text or "").strip() == "JOCKEY":
            node.text = client
            # El modelo tiene "JOCKEY" + espacio + "SALUD"; se limpian esos
            # fragmentos para no dejar restos del cliente ejemplo.
            for cleanup in nodes[index + 1 : index + 4]:
                if (cleanup.text or "").strip() in {"", "SALUD"}:
                    cleanup.text = ""
            return


def _set_cover_line(paragraph: Paragraph, text: str) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.name = "Lucida Sans Unicode"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_TEXT


def update_cover_and_footer(document: Document, payload: dict[str, Any]) -> None:
    client = _clean_text(payload.get("client_name")) or "Cliente"
    period = _clean_text(payload.get("period")) or "Periodo no especificado"
    prepared_by = _clean_text(payload.get("prepared_by")) or "TXDXSECURE"

    for box in document._element.xpath(".//w:txbxContent"):
        original = _textbox_text(box).strip()
        if original == "Change for date":
            _set_textbox_text(box, period)
        elif original == "Change for prepared for":
            _set_textbox_text(box, client)
        elif original == "TXDXSECURE":
            _set_textbox_text(box, prepared_by)
        elif "JOCKEY" in original and "SALUD" in original:
            _replace_textbox_jockey_client(box, client)

    for paragraph in document.paragraphs[:45]:
        text = paragraph.text.strip()
        if "JOCKEY SALUD" in text and "TXDXSECURE" in text:
            _set_cover_line(paragraph, f"{client}\t{prepared_by}")
        elif "Del 20 de junio al 26 de junio del 2026" in text or "Del 20 al 26 de junio del 2026" in text:
            _set_cover_line(paragraph, period)

    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if not footer.paragraphs:
            footer.add_paragraph()
        paragraph = footer.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.tab_stops.clear_all()
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(2.9))
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.15))
        for index, text in enumerate(("TxdxSecure", "Minority Report XOC", client)):
            if index:
                paragraph.add_run("\t")
            run = paragraph.add_run(text)
            run.font.size = Pt(8.5)
            run.font.color.rgb = MUTED_TEXT


def clear_template_body_after_cover(document: Document) -> None:
    """Conserva portada del modelo y remueve índice/contenido ejemplo."""
    body = document._element.body
    children = list(body)
    preserve_until = -1
    for index, child in enumerate(children):
        text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
        if text == "Contenido" or text.startswith("Contenido"):
            preserve_until = index - 1
            while preserve_until >= 0:
                previous = children[preserve_until]
                has_section_break = bool(previous.findall(".//" + qn("w:sectPr")))
                text_before = "".join(node.text or "" for node in previous.iter(qn("w:t"))).strip()
                if has_section_break or text_before:
                    break
                preserve_until -= 1
            break
    if preserve_until < 0:
        last_drawing_index = -1
        for index, child in enumerate(children):
            if child.findall(".//" + qn("w:drawing")):
                last_drawing_index = index
        preserve_until = last_drawing_index if last_drawing_index >= 0 else -1
    preserved = set(children[: preserve_until + 1])
    for child in list(body):
        if child not in preserved and child.tag != "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr":
            body.remove(child)


def _style_name(document: Document, *names: str) -> str:
    for name in names:
        if name in document.styles:
            return name
    return names[-1]


def _format_run(
    run: Any,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    size: float | None = None,
) -> None:
    run.bold = bold or None
    run.italic = italic or None
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)


def apply_example_body_style(document: Document) -> None:
    """Ajusta el cuerpo para parecerse al DOCX ejemplo, usando la portada de la plantilla."""
    styles = document.styles
    if "Body Text" not in styles:
        styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
    if "Normal" in styles:
        normal = styles["Normal"].font
        normal.name = "Tahoma"
        normal.size = Pt(10)
        normal.color.rgb = RGBColor(0x00, 0x00, 0x00)
    if "Heading 1" in styles:
        font = styles["Heading 1"].font
        font.name = "Cambria"
        font.size = Pt(16)
        font.bold = True
        font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    if "Heading 2" in styles:
        font = styles["Heading 2"].font
        font.name = "Arial Black"
        font.size = Pt(12)
        font.bold = True
        font.color.rgb = MODEL_BLUE
    if "Heading 3" in styles:
        font = styles["Heading 3"].font
        font.name = "Arial Black"
        font.size = Pt(10.5)
        font.bold = True
        font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    if "Body Text" in styles:
        body = styles["Body Text"]
        body.font.name = "Tahoma"
        body.font.size = Pt(10)
        body.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body.paragraph_format.space_after = Pt(6)
    if "List Paragraph" in styles:
        list_style = styles["List Paragraph"]
        list_style.font.name = "Tahoma"
        list_style.font.size = Pt(10)
        list_style.paragraph_format.left_indent = Inches(0.25)

    for section in document.sections:
        section.left_margin = Inches(0.64)
        section.right_margin = Inches(0.35)
        section.top_margin = Inches(0.96)
        section.bottom_margin = Inches(0.14)


def _spacing(paragraph: Any, *, before: float = 0, after: float = 7, left: float = 0, first: float | None = None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.left_indent = Pt(left) if left else None
    if first is not None:
        paragraph.paragraph_format.first_line_indent = Pt(first)


def _set_paragraph_bottom_border(paragraph: Any, color: str = "00AEEF", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_heading(document: Document, title: str, number: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = _style_name(document, "Heading 2", "Normal")
    _spacing(paragraph, before=12, after=6)
    run = paragraph.add_run(title)
    _format_run(run, bold=False, color=MODEL_BLUE)


def add_subheading(document: Document, title: str, number: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = _style_name(document, "Heading 3", "Normal")
    _spacing(paragraph, before=7, after=4)
    run = paragraph.add_run(title)
    _format_run(run, bold=False)


def add_body(document: Document, text: Any, *, indent: bool = False) -> None:
    if isinstance(text, list):
        chunks = [_clean_text(item) for item in text if _clean_text(item)]
    else:
        chunks = [chunk.strip() for chunk in str(text or "").splitlines() if chunk.strip()]
    for chunk in chunks:
        paragraph = document.add_paragraph()
        paragraph.style = _style_name(document, "Body Text", "Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _spacing(paragraph, after=5, left=18 if indent else 0)
        run = paragraph.add_run(chunk)


def add_bullets(document: Document, values: Any) -> None:
    for value in _as_list(values):
        text = _clean_text(value)
        if not text:
            continue
        paragraph = document.add_paragraph()
        paragraph.style = _style_name(document, "List Paragraph", "Body Text", "Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _spacing(paragraph, after=4)
        run = paragraph.add_run(text)


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_border(cell: Any, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _set_cell_text(cell: Any, text: Any, *, bold: bool = False, color: RGBColor | None = None, size: float = 9.2) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _spacing(paragraph, after=2)
    run = paragraph.add_run(str(text or ""))
    _format_run(run, bold=bold, color=color, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_callout(document: Document, title: str, content: Any = "", *, fill: str = CALLOUT, bullets: Any = None) -> None:
    text = _clean_text(content)
    bullet_values = [_clean_text(item) for item in _as_list(bullets) if _clean_text(item)] if bullets is not None else []
    if not text and not bullet_values:
        return
    add_subheading(document, title)
    if text:
        add_body(document, text)
    if bullet_values:
        add_bullets(document, bullet_values)


def add_key_value_table(document: Document, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    try:
        table.style = "Table Normal"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(table.rows[0].cells, ("Campo", "Detalle")):
        _shade_cell(cell, "006D9F")
        _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for index, (key, value) in enumerate(rows, start=1):
        cells = table.add_row().cells
        for cell in cells:
            _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cells[0], key, bold=True, color=DARK_TEXT)
        _set_cell_text(cells[1], value)
    document.add_paragraph()


def add_findings_table(document: Document, findings: list[dict[str, Any]]) -> None:
    if not findings:
        return
    table = document.add_table(rows=1, cols=4)
    try:
        table.style = "Table Normal"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(table.rows[0].cells, ("ID", "Vulnerabilidad", "Hosts Afectados", "Severidad")):
        _shade_cell(cell, "006D9F")
        _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8.8)
    for index, finding in enumerate(findings, start=1):
        cells = table.add_row().cells
        for cell in cells:
            _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cells[0], finding.get("id"), bold=False, color=DARK_TEXT, size=8.5)
        _set_cell_text(cells[1], finding.get("vulnerability"), size=8.5)
        _set_cell_text(cells[2], finding.get("affected_hosts"), size=8.5)
        _set_cell_text(cells[3], finding.get("severity"), bold=True, size=8.5)
    document.add_paragraph()


def add_severity_comparison_table(document: Document, rows: list[dict[str, Any]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=3)
    try:
        table.style = "Table Normal"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(table.rows[0].cells, ("Severidad", "Semana Anterior", "Semana Actual")):
        _shade_cell(cell, "006D9F")
        _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8.8)
    for row in rows:
        cells = table.add_row().cells
        for cell in cells:
            _set_cell_border(cell, "BFBFBF", "4")
        _set_cell_text(cells[0], row.get("severity"), size=8.5)
        _set_cell_text(cells[1], row.get("previous"), size=8.5)
        _set_cell_text(cells[2], row.get("current"), size=8.5)
    document.add_paragraph()


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(paragraph, before=2, after=8)
    run = paragraph.add_run(text)
    _format_run(run, italic=True, color=MUTED_TEXT, size=8.5)


def enable_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_word_toc_field(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = _style_name(document, "Normal")

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    paragraph.add_run()._r.append(instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)

    placeholder = paragraph.add_run(
        "Índice automático. Al abrir el documento en Word, actualice los campos si no se muestran las páginas."
    )
    placeholder.italic = True
    placeholder.font.size = Pt(9)
    placeholder.font.color.rgb = MUTED_TEXT

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def document_already_has_page_break(document: Document) -> bool:
    for paragraph in document.paragraphs:
        for br in paragraph._p.findall(".//" + qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
    return False


def add_content_overview(document: Document, payload: dict[str, Any]) -> None:
    if not document_already_has_page_break(document):
        document.add_page_break()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.style = _style_name(document, "Heading 1", "Normal")
    _spacing(title, before=8, after=10)
    run = title.add_run("Contenido")
    _format_run(run, bold=False)

    add_word_toc_field(document)
    document.add_page_break()


def build_report_body(document: Document, payload: dict[str, Any]) -> None:
    enable_update_fields_on_open(document)
    if not any(paragraph.text.strip() == "Contenido" for paragraph in document.paragraphs):
        add_content_overview(document, payload)

    add_heading(document, "Datos generales", "1.")
    add_subheading(document, "Servicio de Monitoreo", "1.1.")
    service_text = (
        f"{payload.get('service_name') or 'Servicio de monitoreo proactivo XOC'} implementado por "
        f"{payload.get('prepared_by') or 'TXDXSECURE'} para el cliente {payload.get('client_name') or 'Cliente'}."
    )
    add_body(document, service_text)
    add_subheading(document, "Periodo", "1.2.")
    add_body(document, payload.get("period"))
    add_subheading(document, "Herramientas", "1.3.")
    tools = payload.get("tools") or []
    if tools:
        add_bullets(document, [f"{tool.get('name')}: {tool.get('description')}" for tool in tools])
    else:
        add_bullets(document, ["No se confirmaron herramientas específicas desde la evidencia entregada."])
    add_subheading(document, "Datos Base", "1.4.")
    add_body(document, payload.get("data_base"))

    add_heading(document, "Resumen ejecutivo del dominio", "2.")
    add_body(document, payload.get("executive_summary"))
    comparison = payload.get("vulnerability_comparison") or {}
    add_subheading(document, "Análisis Comparativo de Vulnerabilidades Semanales", "2.1.")
    add_body(document, comparison.get("summary"))
    severity_rows = comparison.get("severity_rows") or []
    if severity_rows:
        add_severity_comparison_table(document, severity_rows)
    add_subheading(document, "Histograma de la seguridad", "2.2.")
    add_body(document, payload.get("histogram_summary"))
    add_subheading(document, "Resultados obtenidos y próximas acciones", "2.3.")
    add_body(document, payload.get("results_and_next_actions"))
    add_subheading(document, "Resultados obtenidos", "2.4.")
    add_body(document, payload.get("results_obtained"))
    add_subheading(document, "Próximas acciones", "2.5.")
    add_bullets(document, payload.get("next_actions"))
    add_subheading(document, "Requerimiento", "2.5.1.")
    add_bullets(document, payload.get("requirements"))

    add_heading(document, "Seguridad por Dominio", "3.")
    domains = payload.get("security_domains") or []
    for index, domain in enumerate(domains, start=1):
        add_subheading(document, domain.get("name") or f"Dominio {index}", f"3.{index}.")
        add_body(document, domain.get("summary"))
        add_findings_table(document, domain.get("findings") or [])

    add_heading(document, "Reporte de acciones trabajadas durante la semana", "4.")
    add_bullets(document, payload.get("weekly_actions"))

    add_heading(document, "Resultados obtenidos", "5.")
    add_subheading(document, "Seguridad Reforzada", "5.1.")
    add_body(document, payload.get("reinforced_security"))
    add_subheading(document, "Hallazgos pendientes", "5.2.")
    add_bullets(document, payload.get("pending_findings"))

    add_heading(document, "Noticias de seguridad", "6.")
    for index, news in enumerate(payload.get("security_news") or [], start=1):
        add_subheading(document, news.get("title") or f"Noticia {index}", f"6.{index}.")
        add_key_value_table(
            document,
            [
                ("Fecha", news.get("date")),
                ("Fuente", news.get("source")),
                ("Enlaces", ", ".join(news.get("links") or [])),
            ],
        )
        add_body(document, news.get("summary"))
        add_callout(document, "Recomendación", news.get("recommendation"))

    if payload.get("limitations"):
        add_callout(document, "Limitaciones del análisis", bullets=payload.get("limitations"), fill=WARNING)


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _caption_after(paragraph: Paragraph, text: str) -> Paragraph:
    caption = _insert_paragraph_after(paragraph)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED_TEXT
    return caption


def place_evidence_images(document: Document, images: list[Path], descriptions: list[str]) -> None:
    if not images:
        return
    fallback = None
    for paragraph in document.paragraphs:
        if "Resumen ejecutivo" in paragraph.text or "Seguridad por Dominio" in paragraph.text:
            fallback = paragraph
            break
    fallback = fallback or document.paragraphs[-1]
    inserted_after = fallback
    last_citation_anchor_el = None
    for index, path in enumerate(images, start=1):
        label = f"Figura {index}"
        target = None
        for paragraph in document.paragraphs:
            if label.lower() in paragraph.text.lower():
                target = paragraph
                break
        citation_anchor_el = target._p if target is not None else None
        if target is None or citation_anchor_el is last_citation_anchor_el:
            target = inserted_after
        if citation_anchor_el is not None:
            last_citation_anchor_el = citation_anchor_el
        tag = _insert_paragraph_after(target)
        tag.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tag.paragraph_format.space_before = Pt(8)
        tag.paragraph_format.space_after = Pt(3)
        run = tag.add_run(label.upper())
        _format_run(run, bold=True, color=BRAND_BLUE, size=8)
        picture = _insert_paragraph_after(tag)
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(str(path), width=Inches(5.8))
        description = descriptions[index - 1] if index - 1 < len(descriptions) else ""
        caption = _caption_after(picture, f"{label}. {description or 'Evidencia visual proporcionada.'}")
        inserted_after = caption


def validate_docx(path: Path) -> None:
    if not path.exists() or path.stat().st_size < 10_000:
        raise RuntimeError("El DOCX no fue generado correctamente")
    with zipfile.ZipFile(path) as archive:
        if "word/document.xml" not in archive.namelist():
            raise RuntimeError("El DOCX generado no contiene word/document.xml")


def _read_optional_text(path: Path | None) -> str:
    if not path:
        return ""
    return path.read_text(encoding="utf-8")


def _read_optional_json(path: Path | None) -> dict[str, Any]:
    if not path:
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("El JSON estructurado debe ser un objeto.")
    return data


def _text_from_args(args: argparse.Namespace) -> str:
    pieces = []
    if args.text:
        pieces.append(args.text)
    if args.text_file:
        pieces.append(args.text_file.read_text(encoding="utf-8"))
    return "\n\n".join(piece.strip() for piece in pieces if piece.strip())


def main() -> None:
    parser = argparse.ArgumentParser(description="Genera Minority Report XOC usando Azure Foundry multimodal.")
    parser.add_argument("--client-name", required=True, help="Cliente preparado para portada y footer.")
    parser.add_argument("--period", required=True, help="Periodo del reporte. Ejemplo: Del 20 al 26 de junio de 2026.")
    parser.add_argument("--text", default="", help="Texto libre del analista.")
    parser.add_argument("--text-file", type=Path, help="Archivo TXT/MD con contexto del analista.")
    parser.add_argument("--data", type=Path, help="JSON estructurado opcional.")
    parser.add_argument("--reference-md", type=Path, help="Markdown de referencia del Minority Report ejemplo.")
    parser.add_argument("--image", action="append", default=[], type=Path, help="Ruta de imagen local. Se puede repetir.")
    parser.add_argument("--image-description", action="append", default=[], help="Descripción de imagen. Repetir en el mismo orden que --image.")
    parser.add_argument("--no-azure", action="store_true", help="No llama Azure; genera borrador local.")
    parser.add_argument("--allow-local-fallback", action="store_true", default=True, help="Si Azure falla, genera borrador local.")
    args = parser.parse_args()
    load_dotenv(ROOT / ".env")

    try:
        model_path = base_template_path()
    except FileNotFoundError as exc:
        raise SystemExit(str(exc)) from exc

    text = _text_from_args(args)
    structured = _read_optional_json(args.data)
    reference_markdown = _read_optional_text(args.reference_md)
    payload = generate_minority_payload(
        client_name=args.client_name,
        period=args.period,
        analyst_text=text,
        image_paths=args.image,
        image_descriptions=args.image_description,
        structured_data=structured,
        reference_markdown=reference_markdown,
        use_azure=not args.no_azure,
        allow_local_fallback=args.allow_local_fallback,
        output_path=DEFAULT_PAYLOAD_PATH,
    )

    document = Document(model_path)
    apply_example_body_style(document)
    update_cover_and_footer(document, payload)
    clear_template_body_after_cover(document)
    build_report_body(document, payload)
    place_evidence_images(document, args.image, args.image_description)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    result = _writable_output_path(_output_path(payload))
    document.save(result)
    validate_docx(result)
    print(f"Payload JSON: {DEFAULT_PAYLOAD_PATH}")
    print(f"Minority Report generado: {result}")


if __name__ == "__main__":
    main()
